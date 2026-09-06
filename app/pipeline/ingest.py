"""Turn an uploaded file into page images and, where possible, text.

Three paths:
  PDF with a text layer -> read the text directly, no rasterising, no OCR.
  PDF without one       -> render each page at 300 DPI, deskew, contrast.
  .docx                 -> read the text directly.
  image                 -> normalise the single page.
"""
from __future__ import annotations

import hashlib
import os
import re
from pathlib import Path

import cv2
import numpy as np
from PIL import Image, ImageOps

from . import paths, state
from .config import env_bool, env_int
from .imaging import normalize
from .log import get_logger

log = get_logger("ingest")

try:
    import pillow_heif
    pillow_heif.register_heif_opener()
except Exception:  # pragma: no cover
    pass

Image.MAX_IMAGE_PIXELS = 500_000_000


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


# The tags a camera writes and a scanner, a screenshot and an export do not.
# Make and Model sit in the main directory; the rest live in the Exif sub-IFD.
_CAMERA_TAGS_MAIN = (271, 272)                                   # Make, Model
_CAMERA_TAGS_EXIF = (33434, 33437, 34855, 36867, 37386, 42036)   # exposure,
# aperture, ISO, original date, focal length, lens


def looks_photographed(src: Path) -> bool:
    """Whether a camera made this file - asked of the file, not guessed.

    This decides which of the two normalisation paths a page takes, and it
    used to be guessed: an uploaded image file was called a photograph because
    a scan saved as a JPEG and a phone picture of the same page are the same
    file type. That reasoning held only as long as the photograph path could
    not hurt a flat page, and it can. Page-finding is a test of shape, and a
    ruled box printed on a form is the same shape as a sheet of paper; a scan
    that went down the photograph path could be cropped to its own table,
    losing the letterhead above it and the footer below, with what remained
    still reading well enough that OCR raised no objection.

    So it is asked rather than assumed. A phone writes its make, its model and
    its exposure into the file; a flatbed, a screenshot and a PDF export do
    not. A photograph whose tags have been stripped is read as a scan and
    simply gets the treatment it got before any of this existed, which is the
    safe way round: the cost of guessing "scan" is a crop not made, and the
    cost of guessing "photograph" is a page cut in half.
    """
    try:
        with Image.open(src) as im:
            exif = im.getexif()
            if not exif:
                return False
            if any(str(exif.get(tag) or "").strip() for tag in _CAMERA_TAGS_MAIN):
                return True
            sub = exif.get_ifd(0x8769)
            return any(sub.get(tag) not in (None, "") for tag in _CAMERA_TAGS_EXIF)
    except Exception as exc:                     # unreadable metadata is not fatal
        log.warning("%s: could not read the image metadata (%s); "
                    "treating it as a scan", src.name, exc)
        return False


def load_upright(src: Path) -> np.ndarray:
    """Decode an image file the way it was meant to be looked at.

    A phone does not rotate the pixels it saves. It writes them in sensor order
    and records the rotation as an EXIF tag, and every viewer applies that tag
    on the way to the screen - which is why the photograph looks right in the
    gallery and arrives here on its side. Pillow does not apply it either
    unless asked, and OpenCV cannot: it never sees the metadata.

    Nothing downstream recovers from this. Tesseract runs with --psm 3, which
    has no orientation detection, so a sideways page reads as noise and routes
    to the vision model; the vision model is then handed a sideways page and
    answers anyway, in fluent invented text. One tag, applied here, or every
    photograph in the system is read at ninety degrees.
    """
    with Image.open(src) as im:
        upright = ImageOps.exif_transpose(im) or im
        if upright.size != im.size:
            log.info("%s: applied the EXIF rotation (%s -> %s)", src.name,
                     "x".join(map(str, im.size)), "x".join(map(str, upright.size)))
        return cv2.cvtColor(np.array(upright.convert("RGB")), cv2.COLOR_RGB2BGR)


def write_page_image(doc_id: str, page_num: int, image: np.ndarray, *,
                     photo: bool = False) -> Path:
    out, meta = normalize(image, photo=photo)
    if photo:
        # Worth a line each time: page-finding either fired or did not, and
        # which one it was is the first thing to know when a transcript comes
        # back wrong.
        log.info("%s p%s: %s, %dx%d, deskew %.2f deg", doc_id, page_num,
                 ("page found and flattened from " + meta["cropped_from"]
                  if meta.get("page_found") else "no page outline found, "
                  "keeping the whole frame"),
                 meta["width"], meta["height"], meta["deskew_deg"])
    dest = paths.page_image(doc_id, page_num)
    dest.parent.mkdir(parents=True, exist_ok=True)
    # The temp name keeps a .png extension: OpenCV picks its encoder from the
    # extension, and a ".partial" suffix leaves it with no writer at all.
    tmp = dest.with_name(f"{dest.stem}.partial.png")
    if not cv2.imwrite(str(tmp), out, [cv2.IMWRITE_PNG_COMPRESSION, 6]):
        raise RuntimeError(f"could not write {dest.name}")
    os.replace(tmp, dest)
    return dest


def write_text(doc_id: str, page_num: int, text: str) -> Path:
    dest = paths.transcript_txt(doc_id, page_num)
    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_text(text.rstrip() + "\n", encoding="utf-8")
    return dest


def _record_page(conn, doc_id: str, page_num: int, *, image: Path | None = None,
                 text: Path | None = None, source: str | None = None,
                 route: str | None = None) -> None:
    conn.execute(
        """INSERT INTO pages (doc_id, page_num, image_path, text_path, text_source, route)
           VALUES (?,?,?,?,?,?)
           ON CONFLICT(doc_id, page_num) DO UPDATE SET
             image_path  = COALESCE(excluded.image_path, pages.image_path),
             text_path   = COALESCE(excluded.text_path, pages.text_path),
             text_source = COALESCE(excluded.text_source, pages.text_source),
             route       = COALESCE(excluded.route, pages.route)""",
        (doc_id, page_num,
         paths.rel(image) if image else None,
         paths.rel(text) if text else None, source, route))


def ingest_pdf(src: Path, doc_id: str, on_progress) -> int:
    from pdf2image import convert_from_path
    from pypdf import PdfReader

    reader = PdfReader(str(src))
    page_count = len(reader.pages)
    if page_count == 0:
        raise RuntimeError("PDF reports zero pages")

    min_chars = env_int("EMBEDDED_TEXT_MIN_CHARS", 120)
    use_layer = env_bool("USE_EMBEDDED_TEXT_LAYER", True)

    for page_num in range(1, page_count + 1):
        on_progress(f"reading page {page_num}/{page_count}")
        text = ""
        if use_layer:
            try:
                text = (reader.pages[page_num - 1].extract_text() or "").strip()
            except Exception:
                text = ""

        if len(text) >= min_chars:
            # Born-digital page: rasterising and re-reading it would be slower
            # and strictly lossier than the text already in the file.
            txt = write_text(doc_id, page_num, text)
            with state.tx() as conn:
                _record_page(conn, doc_id, page_num, text=txt,
                             source="pdf_text", route="text")
            continue

        if not paths.page_image(doc_id, page_num).exists():
            # One page at a time: a long PDF at 300 DPI will not fit in memory.
            images = convert_from_path(str(src), dpi=env_int("PDF_DPI", 300),
                                       first_page=page_num, last_page=page_num,
                                       fmt="png")
            if not images:
                raise RuntimeError(f"page {page_num} produced no image")
            arr = cv2.cvtColor(np.array(images[0].convert("RGB")), cv2.COLOR_RGB2BGR)
            write_page_image(doc_id, page_num, arr)
            del images, arr
        with state.tx() as conn:
            _record_page(conn, doc_id, page_num,
                         image=paths.page_image(doc_id, page_num))
    return page_count


def ingest_docx(src: Path, doc_id: str, on_progress) -> int:
    """Word text, including tables, which often carry the actual details."""
    from docx import Document as DocxDocument

    on_progress("reading document text")
    document = DocxDocument(str(src))
    parts: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    text = "\n".join(line for line in parts if line.strip())
    if not text.strip():
        raise RuntimeError("no readable text in this Word document")

    txt = write_text(doc_id, 1, text)
    with state.tx() as conn:
        _record_page(conn, doc_id, 1, text=txt, source="docx", route="text")
    return 1


def ingest_image(src: Path, doc_id: str, on_progress) -> int:
    on_progress("normalising image")
    arr = load_upright(src)
    photo = looks_photographed(src)
    log.info("%s: %s", src.name,
             "a camera made this file, so it takes the photograph path"
             if photo else "no camera metadata, so it is normalised as a scan")
    dest = write_page_image(doc_id, 1, arr, photo=photo)
    with state.tx() as conn:
        _record_page(conn, doc_id, 1, image=dest)
    return 1


# What a document IS decides how much its contents are worth. A system log is
# the record; a person describing that log from memory is a restatement of it.
# Citing the restatement when the record is in evidence is the sourcing error
# that survives every prose instruction, so the kind is stored as data.
# Order matters more than the patterns do. A transcript that discusses a log is
# still a transcript, so the structural kinds - anything with questions and
# answers, a sworn preamble, a memorandum heading - are tested BEFORE the
# content words that suggest a record. Testing record first classified an
# interview as a log because the witness talked about logs.
_KIND_PATTERNS = [
    ("statement", re.compile(
        r"\b(sworn statement|statement of witness|affidavit|"
        r"under (oath|penalty of perjury))", re.I)),
    ("interview", re.compile(
        r"(^|\n)\s*(q|question|io)\s*[:.]|\b(interview (of|with|transcript)|"
        r"transcript of interview|interviewee\s*[:.])", re.I)),
    ("appointment", re.compile(
        r"\b(memorandum for|appointment (memo|letter)|"
        r"appointed to conduct|investigating officer is appointed)", re.I)),
    ("notes", re.compile(
        r"\b(working notes|io notes|investigator(?:'s)? notes)", re.I)),
    ("record", re.compile(
        r"\b(log|logbook|ledger|register|system report|report generated|"
        r"proxy log|access log|audit trail|export|certificate|"
        r"calibration record|maintenance record|orders|form \d)", re.I)),
]


def classify_kind(filename: str, first_text: str) -> str:
    """Best-effort document kind from the filename and its opening text."""
    blob = f"{filename}\n{first_text[:1200]}"
    for kind, pattern in _KIND_PATTERNS:
        if pattern.search(blob):
            return kind
    return "unknown"


# Who is speaking decides how much weight their account carries about a given
# thing. A custodian describing the system they administer is the source for
# what that system recorded; the subject repeating the same figure from memory
# is a restatement of it. Both are interviews, so document kind cannot tell
# them apart - the speaker's relationship to the evidence has to be captured.
_ROLE_PATTERNS = [
    ("subject", re.compile(
        r"\b(subject of (this|the) investigation|you are the subject|"
        r"subject interview|as the subject)\b", re.I)),
    ("custodian", re.compile(
        r"\b(custodian|network administrator|system administrator|"
        r"records? (manager|monitor|keeper)|calibration monitor|"
        r"i (maintain|administer|keep|run) the (log|system|records?)|"
        r"i pulled the (log|report)|i am responsible for the records?)\b", re.I)),
    ("complainant", re.compile(
        r"\b(complainant|i filed (a|the) complaint|i reported (him|her|them|it) to)\b",
        re.I)),
    ("supervisor", re.compile(
        r"\b(section chief|flight chief|supervisor|i supervise|"
        r"in my capacity as (his|her|their) supervisor)\b", re.I)),
]


def classify_role(filename: str, first_text: str) -> str:
    """The interviewee's relationship to the evidence, best effort.

    Text is weighted over the filename: a file named for a role is a
    convention this pipeline cannot rely on, but a person saying what they do
    is present in any real interview.
    """
    for role, pattern in _ROLE_PATTERNS:
        if pattern.search(first_text[:2500]):
            return role
    stem = filename.lower()
    for role, hints in (("subject", ("subject",)),
                        ("custodian", ("custodian", "admin", "monitor", "records")),
                        ("complainant", ("complainant",)),
                        ("supervisor", ("supervisor", "sectionchief", "chief"))):
        if any(h in stem for h in hints):
            return role
    return "witness"


def run(doc_id: str, on_progress) -> int:
    doc = state.query_one("SELECT * FROM documents WHERE doc_id=?", (doc_id,))
    src = paths.RAW / doc["filename"]
    suffix = src.suffix.lower()

    if suffix in paths.SUPPORTED_PDF:
        page_count = ingest_pdf(src, doc_id, on_progress)
    elif suffix in paths.SUPPORTED_DOC:
        page_count = ingest_docx(src, doc_id, on_progress)
    else:
        page_count = ingest_image(src, doc_id, on_progress)

    first_text = ""
    first = paths.transcript_txt(doc_id, 1)
    if first.exists():
        first_text = first.read_text(encoding="utf-8")[:1200]
    kind = classify_kind(doc["filename"], first_text)
    role = classify_role(doc["filename"], first_text) if kind in (
        "interview", "statement", "unknown") else ""

    with state.tx() as conn:
        conn.execute(
            "UPDATE documents SET page_count=?, doc_kind=?, doc_role=? WHERE doc_id=?",
            (page_count, kind, role, doc_id))
    log.info("%s: %s%s", doc_id, kind, f" / {role}" if role else "")
    log.info("%s: %d page(s)", doc_id, page_count)
    return page_count
