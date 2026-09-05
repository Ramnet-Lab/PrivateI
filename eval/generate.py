#!/usr/bin/env python3
"""Generate a synthetic investigation case that plants one instance of every
failure mechanic, plus the machine-readable key to grade against.

Seeded: names, units, buildings, and dates all vary with --seed, so a pipeline
change that improves the score is improving the mechanics - not memorising one
fixture. Anchor tokens (serial numbers, figures) are unique per seed so the
graders can match on strings the model cannot paraphrase away.
"""
from __future__ import annotations

import argparse
import json
import random
from datetime import date, timedelta
from pathlib import Path

from docx import Document

RANKS = ["SSgt", "TSgt", "MSgt", "SrA", "A1C", "Capt"]
FIRST = ["Jordan", "Casey", "Riley", "Avery", "Dana", "Quinn", "Morgan",
         "Devin", "Skyler", "Harper", "Rowan", "Blake"]
LAST = ["Calloway", "Mercer", "Ibarra", "Whitfield", "Okonkwo", "Lindqvist",
        "Marsh", "Trevino", "Halvorsen", "Bright", "Suzuki", "Delacroix"]
SHOPS = ["Fuels Flight", "Avionics Backshop", "Munitions Storage",
         "Vehicle Maintenance", "Comm Focal Point"]


def build(seed: int, out: Path) -> dict:
    rng = random.Random(seed)
    out.mkdir(parents=True, exist_ok=True)

    last = rng.sample(LAST, 6)
    first = rng.sample(FIRST, 6)
    subj = f"MSgt {first[0]} {last[0]}"          # the subject
    w_a = f"SSgt {first[1]} {last[1]}"           # eyewitness A
    w_b = f"SrA {first[2]} {last[2]}"            # eyewitness B, limited view
    w_c = f"A1C {first[3]} {last[3]}"            # complainant
    custodian = f"Ms. {first[4]} {last[4]}"      # records custodian
    decoy = f"SrA {first[2]} {last[5]}"          # same FIRST name as w_b, different person
    shop = rng.choice(SHOPS)
    bldg = f"Building {rng.randint(200, 980)}"
    serial = f"TL-{rng.randint(1000, 9999)}"     # anchor: a tool serial
    figure = rng.randint(800, 4000)              # anchor: a records figure
    base_day = date(2026, rng.randint(3, 6), rng.randint(3, 20))
    incident = base_day + timedelta(days=10)
    boundary_day = base_day + timedelta(days=3)
    anchor_day = incident - timedelta(days=incident.weekday())  # a Monday <= incident

    def d(x): return x.strftime("%-d %b %Y") if hasattr(x, "strftime") else x
    def iso(x): return x.isoformat()

    def write_docx(name: str, title: str, meta: list[str], paras: list[str]):
        doc = Document()
        doc.add_heading(title, 1)
        for m in meta:
            doc.add_paragraph(m)
        doc.add_paragraph("")
        for p in paras:
            doc.add_paragraph(p)
        doc.save(out / name)

    # --- 01 appointment: goal + numbered allegations (alignment mechanic) ---
    write_docx("01_appointment.docx", "APPOINTMENT OF INVESTIGATING OFFICER",
        ["MEMORANDUM FOR THE INVESTIGATING OFFICER"],
        [f"The goal of this investigation is to determine the facts and "
         f"circumstances surrounding alleged misconduct by {subj}, NCOIC, {shop}.",
         f"Allegation 1: That on {d(incident)}, {subj} made demeaning remarks "
         f"to subordinates in the {shop} tool room.",
         f"Allegation 2: That {subj} failed to perform weekly tool "
         f"accountability checks after {d(boundary_day)}.",
         f"Allegation 3: That {subj} removed tool {serial} from {bldg} "
         f"for personal use."])

    # --- 02 witness A: negation + wording variance + kin ---
    write_docx("02_statement_A.docx", "SWORN STATEMENT",
        [f"Name: {w_a}", f"Organization: {shop}", f"Date: {d(incident + timedelta(days=6))}"],
        [f"I was in the tool room on {d(incident)}